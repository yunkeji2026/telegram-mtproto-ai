#!/usr/bin/env python3
"""生成《系统能力说明与造价评估》Word 文档。依赖: pip install python-docx"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_cell_shading(cell, fill_hex: str) -> None:
    from docx.oxml import parse_xml
    shading = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
            fill_hex,
        )
    )
    cell._tc.get_or_add_tcPr().append(shading)


def main() -> None:
    out = Path(__file__).resolve().parent / "系统能力说明与造价评估.docx"
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # 默认字体：中文宋体/西文 Times 由 Word 接管；此处设正文基础
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _style_title_run(run, size_pt: int, color: RGBColor, bold: bool = True) -> None:
        run.bold = bold
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        run.font.name = "Microsoft YaHei"
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), "微软雅黑")

    # 标题页
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Telegram MTProto AI 客服系统")
    _style_title_run(r, 22, RGBColor(0x1A, 0x3C, 0x6E))

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("能力说明与造价评估（务实版）")
    _style_title_run(r2, 16, RGBColor(0x44, 0x44, 0x44))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("（基于当前代码规模与常见国内交付行情）")
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    s.font.name = "Microsoft YaHei"
    s._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = p.add_run("内部资料 · 仅供决策参考")
    pr.italic = True
    pr.font.size = Pt(9)
    pr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # 一、文档目的
    doc.add_heading("一、文档目的", level=1)
    doc.add_paragraph(
        "本文档用于概括本系统已实现的核心能力，并给出更符合国内市场习惯的「从零开发同等系统」"
        "造价区间，便于对外报价、对内立项或融资材料使用。"
    )

    # 二、系统能力概要
    doc.add_heading("二、系统能力概要", level=1)
    items = [
        "接入方式：基于 Pyrogram 的 Telegram 用户（MTProto）客户端，非仅 Bot API，可处理群内复杂交互。",
        "智能回复：对接 OpenAI 兼容 API（如 DeepSeek 等），支持意图识别、多轮对话、知识库（BM25/向量等）与领域技能插件。",
        "业务域：支付域（通道状态、额度、订单/GXP 等）可扩展；支持多域加载与技能注册。",
        "触发与风控：四层触发决策、群回复模式（@/关键词/回复链等）、收窄回复（narrow_reply）、限流与人工升级策略。",
        "运营侧：Web 管理端（配置、知识库、部分运营能力）、监控与日志等工程化能力。",
        "多模态扩展点：语音/图片/OCR/Vision 等可按配置启用。",
    ]
    for x in items:
        doc.add_paragraph(x, style="List Bullet")

    # 三、规模参考
    doc.add_heading("三、规模与工作量参考", level=1)
    doc.add_paragraph(
        "以本仓库 telegram-mtproto-ai 为例，Python 代码体量约在「数万行」量级（含领域技能、"
        "客户端、触发器、Web、工具与测试等）。该规模在行业中通常对应「多模块、可上线运营」的中型后台产品，"
        "而非演示级脚本。"
    )
    doc.add_paragraph(
        "若从零实现到可生产部署，通常需要「产品/架构 + 后端 + 部分前端 + 联调测试」协作，"
        "综合有效工作量常见落在约 6～12 人月（视需求冻结程度与质量要求浮动）。"
    )

    # 四、为何常见「首轮估价」容易偏高
    doc.add_heading("四、为何首轮估价容易显得偏高", level=1)
    doc.add_paragraph(
        "部分初次估算会按「一线城市软件公司全包交付」口径：含较高人天单价、项目管理与利润、"
        "完整文档与验收、以及隐含风险溢价。该口径对「预算有限、需求已清晰、可接受分阶段交付」"
        "的项目而言，会显得偏贵。"
    )
    doc.add_paragraph(
        "更务实的做法是分场景：内部自建、中小外包团队、成熟软件公司、是否含首年运维，"
        "分别对应不同单价与总价区间。"
    )

    # 五、务实价格评估表
    doc.add_heading("五、市场价格评估（务实区间）", level=1)
    doc.add_paragraph(
        "下表为「从零开发到功能与当前系统大致等价、可上线使用」的国内市场常见议价区间，"
        "货币为人民币，含税与否、是否含第三方费用（云服务器、模型 API 调用费等）需单独约定。"
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "交付场景"
    hdr[1].text = "典型假设"
    hdr[2].text = "总价常见区间（人民币）"
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(6.8)
    table.columns[2].width = Cm(5.5)
    for c in hdr:
        for p in c.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.bold = True
        set_cell_shading(c, "D9E2F3")

    rows_data = [
        (
            "内部/小团队自建",
            "核心成员兼职或低管理成本，需求边界清晰",
            "约 10 万～28 万\n（主要为人力与时间成本，非对外报价）",
        ),
        (
            "二三线城市外包团队",
            "6～8 人月交付，中等文档与测试",
            "约 18 万～40 万",
        ),
        (
            "一线城市成熟外包 / 软件公司",
            "含项目管理、测试与验收文档，风险溢价适中",
            "约 28 万～55 万",
        ),
        (
            "品牌公司或大厂合作部",
            "强流程、审计与 SLA，管理成本高",
            "约 45 万～90 万+\n（非多数项目的必选）",
        ),
        (
            "含首年维护与中等规模二开额度",
            "Bug 修复 + 小功能迭代包",
            "在上表基础上通常 +15%～35%",
        ),
    ]
    for a, b, c in rows_data:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c

    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run("说明：")
    nr.bold = True
    note.add_run(
        "模型 API、短信、服务器与域名等运行成本为「持续运营费用」，一般不计入一次性开发总价，"
        "或按年单独列支。"
    )

    # 六、综合结论（合理落点）
    doc.add_heading("六、综合结论（合理落点）", level=1)
    doc.add_paragraph(
        "在需求范围清晰、采用国内常见外包或小型软件公司交付的前提下，"
        "与本项目复杂度相近的系统，「合理可谈」的开发总价多数会落在："
    )
    concl = doc.add_paragraph()
    cr = concl.add_run("约 人民币 22 万～45 万")
    cr.bold = True
    cr.font.size = Pt(13)
    cr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    concl.add_run(
        "（中间位常接近 28～35 万）。若仅追求可用 MVP、接受文档与测试从简，可下探至约 15 万～22 万区间；"
        "若必须对标一线品牌外包交付，则仍可能显著高于上述中位。"
    )

    doc.add_heading("七、免责声明", level=1)
    doc.add_paragraph(
        "本文为基于公开代码规模与行业经验的估算，不构成正式报价或审计结论。"
        "实际价格随需求变更、验收标准、知识产权归属、付款方式与地域差异而波动，"
        "商务谈判以书面合同为准。"
    )

    doc.save(out)
    print(f"已生成: {out}")


if __name__ == "__main__":
    main()
