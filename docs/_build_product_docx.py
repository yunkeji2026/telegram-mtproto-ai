# -*- coding: utf-8 -*-
"""One-off: build 回复逻辑与知识库面板说明-产品版.docx"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent / "回复逻辑与知识库面板说明-产品版.docx"


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def main():
    doc = Document()
    title = doc.add_heading("智能客服机器人：回复逻辑与后台使用说明（产品版）", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "本文面向产品、运营与非研发同事，说明当前机器人在群里「什么时候说话、说什么」、"
        "后台知识库与面板各自负责什么，以及哪些能力暂时用不上或仅作备用。"
    )

    doc.add_heading("一、机器人能帮你解决什么", level=1)
    add_body(
        doc,
        "当前策略下，机器人主要服务「支付/通道」类群场景：在群里被触发后，"
        "优先回答与客服在线、通道状态、限额与成功率等相关问题；"
        "其它类型问题多数不自动回复，避免误答或泄露敏感信息。"
    )

    doc.add_heading("二、用户侧：什么情况下会回复", level=1)
    doc.add_paragraph("会回复的典型情况（需同时满足群内触发规则，例如关键词、@ 机器人等，具体以实际配置为准）：", style="List Bullet")
    doc.add_paragraph("用户问客服是否在、寒暄（如你好、在吗、单独一个「在」等）。", style="List Bullet")
    doc.add_paragraph("用户咨询通道是否可用、额度/限额、成功率、是否维护、代收代付状态等。", style="List Bullet")
    doc.add_paragraph("在上一条刚聊过通道/限额/成功率的情况下，用户短句追问（如「正常吗」「？」）。", style="List Bullet")

    doc.add_paragraph("一般不会回复或会被拦截的情况：", style="List Bullet")
    doc.add_paragraph("涉及订单号、查单、投诉、退款、凭证、到账异常等——需要人工或专用流程处理。", style="List Bullet")
    doc.add_paragraph("被系统识别为闲聊、订单查询等、但不在当前「收窄」允许范围内时，机器人可能不回复。", style="List Bullet")

    doc.add_heading("三、什么是「收窄回复」", level=1)
    add_body(
        doc,
        "可以理解为产品策略：在群里自动回复时，只处理少数几类用户意图（例如问候类、通道与状态类），"
        "其余意图不自动回复。这样可减少误触、控制话术风险，并把复杂问题留给人工。"
    )
    add_body(
        doc,
        "若以后业务需要「恢复查单、投诉等自动应答」，需要在配置与运营策略上一起调整，不是单改某一条知识库就能实现。"
    )

    doc.add_heading("四、后台知识库（词条）是干什么的", level=1)
    add_body(
        doc,
        "知识库存的是「标准说法、步骤、注意事项、示例话术」等，供机器人在回答时参考。"
        "用户问法千变万化，系统会用检索把最相关的条目找出来，再组织成自然语言回复。"
    )
    doc.add_paragraph("词条以外，后台还有：", style="List Bullet")
    doc.add_paragraph("错误码说明：用户消息里出现错误码时，辅助解释。", style="List Bullet")
    doc.add_paragraph("规则与对话示例：用于约束回答风格或提供范例，视配置参与上下文。", style="List Bullet")
    add_body(
        doc,
        "特别说明：当用户明确问「成功率、费率、手续费」等时，系统会优先使用「通道实时状态」里的数据来答，"
        "避免知识库里过期的费率文案抢先回答。词条里若写过「去后台查费率」，在收窄策略下也不应作为主要出口。"
    )
    add_body(
        doc,
        "运营维护建议：把词条按分类整理好；禁用或更新过时条目；重要话术与法务/合规对齐后再上线。"
    )

    doc.add_heading("五、管理后台（网页）常用页面", level=1)
    doc.add_paragraph("在浏览器打开本机管理地址（通常为 http://127.0.0.1:8080 ，登录后以实际环境为准）：", style="List Number")
    doc.add_paragraph("首页/仪表盘：总览通道健康、模板数量、近期操作等。", style="List Number")
    doc.add_paragraph("通道管理：维护各支付通道的状态、成功率、限额等——这些会直接影响机器人回答里的「实时通道」信息。", style="List Number")
    doc.add_paragraph("模板：部分固定句式（如问候、兜底），与单条知识库词条是不同入口。", style="List Number")
    doc.add_paragraph("知识库：增删改词条、导入导出、翻译、备份等，是日常运营的主入口。", style="List Number")
    doc.add_paragraph("设置：含回复逻辑、意图关键词等，与配置文件联动，改动前建议备份或找技术同事确认。", style="List Number")
    doc.add_paragraph("日志与审计：排查问题、查看谁在何时改了配置。", style="List Number")
    add_body(
        doc,
        "说明：支付业务专属的独立页面若未单独开发，通道与汇率等仍集中在上述后台中，以实际菜单为准。"
    )

    doc.add_heading("六、哪些能力算「暂时少用」或备用", level=1)
    doc.add_paragraph("在收窄策略下，大量与「查单、投诉、闲聊」相关的知识库条目，当前自动回复链路可能用不到；"
                      "词条可保留作日后策略放宽或培训材料。", style="List Bullet")
    doc.add_paragraph("自动调参、自动观察类开关若处于关闭状态，则不会自动调整策略。", style="List Bullet")
    doc.add_paragraph("向量/语义增强检索依赖 embedding 配置；未配置时以关键词检索为主，一般仍可用。", style="List Bullet")
    doc.add_paragraph("学习型、批量进化、部分高级实验功能，偏运营与进阶场景，非日常必用。", style="List Bullet")

    doc.add_heading("七、与人工客服的配合", level=1)
    add_body(
        doc,
        "机器人不回复不等于用户被忽视：可结合群内人工转接规则、@ 指定客服、工作时间等，"
        "把「收窄」省下的精力集中在高价值或高风险会话上。具体转接话术与排班以《人工转接使用说明》等文档为准。"
    )

    doc.add_paragraph("")
    p = doc.add_paragraph("文档版本：产品版 · 生成于部署仓库，修订时请同步业务现状。")
    p.runs[0].italic = True

    doc.save(OUT)
    print("Wrote:", OUT)


if __name__ == "__main__":
    main()
