# -*- coding: utf-8 -*-
"""生成 Camille 智能客服操作手册 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── 样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.3

BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x1B, 0x20, 0x38)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_BG = "D6E4FF"
TABLE_HEADER_BG = "1A56DB"
TABLE_ALT_BG = "F0F4FF"
ACCENT_GREEN = RGBColor(0x10, 0xB9, 0x81)


def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tcPr.append(shading)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK if level == 1 else BLUE
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h


def add_para(text, bold=False, color=None, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if align:
        p.alignment = align
    return p


def add_tip(text, icon="💡"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{icon} {text}")
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY
    run.italic = True
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, TABLE_HEADER_BG)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if ri % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.name = '微软雅黑'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run2.font.size = Pt(10)
    return p


def add_numbered(text, num):
    p = doc.add_paragraph()
    run_num = p.add_run(f"  {num}  ")
    run_num.bold = True
    run_num.font.color.rgb = BLUE
    run_num.font.size = Pt(11)
    run_num.font.name = '微软雅黑'
    run_num.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    return p


# ═══════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

add_para("Camille 智能客服", bold=True, size=Pt(28), color=DARK,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("操 作 手 册", bold=True, size=Pt(22), color=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para("—— 写给客服团队的使用指南 ——", size=Pt(12), color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(6):
    doc.add_paragraph()

add_para("版本：v2.0  |  2026年3月", size=Pt(10), color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("本手册用最简单的话，教你怎么用后台管理 Camille 客服机器人。",
         size=Pt(10), color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 目录
# ═══════════════════════════════════════════════════

add_heading_styled("目 录", 1)
toc_items = [
    "一、Camille 是什么",
    "二、后台管理入口",
    "三、仪表盘 — 一眼看懂运行状态",
    "四、知识库 — 教 Camille 回答问题",
    "五、智能学习 — 让 Camille 自己成长",
    "六、回复逻辑 — 控制什么时候回复",
    "七、策略管理 — 调整回复风格",
    "八、人工转接 — 搞不定时叫人",
    "九、常用操作速查表",
    "十、遇到问题怎么办",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 一、Camille 是什么
# ═══════════════════════════════════════════════════

add_heading_styled("一、Camille 是什么", 1)

add_para("Camille 是一个智能客服机器人，运行在 Telegram 上。她可以：")
add_bullet("自动回复群里用户的问题（订单、额度、通道状态等）")
add_bullet("识别用户发的图片和语音")
add_bullet("用用户的语言回复（中文、英文、阿拉伯语等都可以）")
add_bullet("遇到搞不定的问题，自动叫人工客服来帮忙")
add_bullet("每天自动学习，发现新问题并生成回答草稿")

add_tip("Camille 不会乱编订单信息。如果用户没发截图或订单号，她只会说「请发订单号或截图」。")

doc.add_paragraph()
add_para("Camille 能处理的问题类型：", bold=True)
add_table(
    ["问题类型", "举例", "Camille 怎么做"],
    [
        ["查订单", "「帮我查一下 12345678」", "引导发单号或截图，有凭证时确认订单状态"],
        ["问额度", "「EP 额度多少」「JC 限额」", "直接回复准确的额度数字"],
        ["问通道", "「成功率怎么样」「稳定吗」", "回复各通道当前状态"],
        ["打招呼", "「你好」「在吗」", "友好问候，问有什么能帮忙的"],
        ["投诉", "「怎么还没到账」「太慢了」", "安抚情绪，引导发凭证核查"],
        ["代发命令", "「查汇率」「查余额」", "帮用户向系统机器人发送查询命令"],
    ],
    [3, 6, 7]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 二、后台管理入口
# ═══════════════════════════════════════════════════

add_heading_styled("二、后台管理入口", 1)

add_para("打开浏览器，在地址栏输入：")
add_para("http://你的服务器IP:8080", bold=True, size=Pt(12), color=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para("首次使用需要创建管理员账号。之后每次用账号密码登录即可。")
add_para("登录后你会看到左侧有一排菜单：", bold=True)

add_table(
    ["菜单", "干什么用"],
    [
        ["仪表盘", "看 Camille 现在运行得怎么样"],
        ["知识库", "管理 Camille 的「问题-回答」库"],
        ["智能学习", "看 Camille 自动发现的新问题和生成的回答"],
        ["策略管理", "调整 Camille 回复的温度、长度"],
        ["策略效果", "看各种回复策略的效果对比"],
        ["系统设置", "改回复逻辑、AI参数、人工转接等"],
        ["操作记录", "看谁在后台做了什么操作"],
        ["实时日志", "看 Camille 的运行日志"],
    ],
    [3.5, 12.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 三、仪表盘
# ═══════════════════════════════════════════════════

add_heading_styled("三、仪表盘 — 一眼看懂运行状态", 1)

add_para("仪表盘是你打开后台第一个看到的页面。上面几个数字告诉你 Camille 现在的状态。")

add_table(
    ["指标", "什么意思", "正常范围"],
    [
        ["Telegram 连接", "Camille 和 Telegram 是否连上了", "显示「已连接」就正常"],
        ["AI 状态", "AI 大脑是否正常工作", "显示「健康」就正常"],
        ["收到消息数", "Camille 今天收到了多少条消息", "正常有数字就行"],
        ["已回复消息", "Camille 今天回复了多少条", "正常有数字就行"],
        ["知识库条目", "Camille 知道多少条问答", "越多越好"],
    ],
    [3.5, 7, 5.5]
)

add_tip("如果 Telegram 连接断了或 AI 不健康，请联系技术人员。")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 四、知识库
# ═══════════════════════════════════════════════════

add_heading_styled("四、知识库 — 教 Camille 回答问题", 1)

add_para("知识库就是 Camille 的「题库」。用户问了一个问题，Camille 会先去知识库里找有没有对应的回答。")

add_heading_styled("怎么添加新的问答", 2)
add_numbered("点左边菜单「知识库」", "①")
add_numbered("点右上角「新增条目」按钮", "②")
add_numbered("填写以下内容：", "③")

add_table(
    ["要填的内容", "怎么填", "举例"],
    [
        ["分类", "选一个类别", "订单查询 / 通道状态 / 常规咨询"],
        ["标题", "简单描述这个问题", "用户问 EP 通道额度"],
        ["触发词", "用户可能怎么问（可以写多个，逗号隔开）", "EP额度, ep限额, EP多少"],
        ["建议回复", "Camille 应该怎么回答", "EP 普通额度为 100-20,000"],
    ],
    [3, 6, 7]
)

add_numbered("点「保存」", "④")

add_tip("触发词越多，用户问到时 Camille 越容易找到这条回答。")

add_heading_styled("怎么修改已有的回答", 2)
add_numbered("在知识库页面找到那条记录", "①")
add_numbered("点「编辑」", "②")
add_numbered("改完后点「保存」", "③")

add_heading_styled("怎么搜索知识库", 2)
add_para("在知识库页面顶部有搜索框，输入关键词就能过滤。")

add_heading_styled("沙盒测试 — 验证效果", 2)
add_para("知识库页面有「沙盒测试」功能。你输入一句模拟用户的话，Camille 会告诉你她会怎么回复。用这个来验证新加的知识条目是否生效。")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 五、智能学习
# ═══════════════════════════════════════════════════

add_heading_styled("五、智能学习 — 让 Camille 自己成长", 1)

add_para("Camille 会自动收集她回答不了的问题，然后用 AI 生成回答草稿。你只需要审核就行。")

add_heading_styled("工作流程", 2)

add_numbered("Camille 每天自动收集「回答不了的问题」和「用户不满意的回答」", "①")
add_numbered("AI 自动为每个问题生成一份「回答草稿」", "②")
add_numbered("你在后台审核：觉得好 → 通过入库；不好 → 拒绝；差不多 → 编辑后通过", "③")

add_heading_styled("智能学习流程图", 2)
img_path2 = os.path.join(os.path.dirname(__file__), "learning_flow.png")
if os.path.exists(img_path2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path2, width=Cm(14))

add_heading_styled("怎么操作", 2)
add_numbered("点左边菜单「智能学习」", "①")
add_numbered("你会看到待审核的草稿列表", "②")
add_numbered("每条草稿显示：", "③")

add_bullet("用户原来问了什么", "问题：")
add_bullet("AI 建议的回答", "建议回复：")
add_bullet("AI 建议的分类和触发词", "分类/触发词：")

add_para("对每条草稿你可以：", bold=True)
add_table(
    ["按钮", "效果"],
    [
        ["通过入库", "这条回答直接加到知识库，Camille 以后就会用这个回答"],
        ["拒绝", "扔掉，不要这条"],
        ["编辑", "改一改再通过"],
        ["全部通过", "一次性把所有待审核的都通过（慎用）"],
    ],
    [4, 12]
)

add_heading_styled("手动触发学习", 2)
add_para("不想等每天自动跑，可以点右上角「立即学习」按钮，马上收集一轮。")

add_tip("建议每天花 5 分钟看一下智能学习页面，通过好的、拒绝差的。Camille 会越来越聪明。")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 六、回复逻辑
# ═══════════════════════════════════════════════════

add_heading_styled("六、回复逻辑 — 控制什么时候回复", 1)

add_para("这是最重要的设置之一。它决定了 Camille 在群里什么消息该回、什么消息不回。")

add_para("位置：系统设置 → 回复逻辑管理", bold=True, color=BLUE)

add_heading_styled("核心开关", 2)

add_table(
    ["开关", "什么意思", "建议"],
    [
        ["四层智能触发", "Camille 自动判断这条消息是不是在问她，只有判断是的才回复",
         "建议开启（默认就是开的）"],
        ["群消息回复模式", "四层触发关掉后的备用模式",
         "建议选「@ 或关键词」"],
    ],
    [4, 7, 5]
)

add_heading_styled("回复逻辑流程图", 2)
img_path = os.path.join(os.path.dirname(__file__), "reply_logic_flow.png")
if os.path.exists(img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Cm(14))

add_heading_styled("四层触发是什么意思（白话版）", 2)
add_table(
    ["层", "Camille 干什么", "举例"],
    [
        ["第1层", "消息里有关键词或订单号 → 立刻回复", "用户发「查单 12345」"],
        ["第2层", "第1层没命中，AI 判断这条消息是不是在问客服", "用户发「今天能跑吗」"],
        ["第3层", "通过了第2层，但检查冷却时间，刚回过的不重复回", "30秒内第2次问同样的话"],
        ["结果", "都没通过 → 不回复，保持安静", "用户在聊天，跟客服无关"],
    ],
    [2, 7, 7]
)

add_heading_styled("子开关说明", 2)
add_table(
    ["子开关", "什么意思", "建议"],
    [
        ["回复链触发", "用户回复 Camille 之前发的消息时，一定回复", "建议开"],
        ["追问检测", "最近回过这个用户，他后续的消息更容易触发回复", "建议开"],
        ["会话窗口", "回复过某用户后的一段时间内（比如45分钟），对他更积极", "建议开"],
        ["L2 兜底", "主触发没命中时，AI 再判一次要不要回", "按需，回复太多可以关"],
        ["AI 上下文判断", "AI 看上下文决定要不要回", "按需，回复太多可以关"],
    ],
    [4, 7, 5]
)

add_heading_styled("两个关键滑块", 2)
add_table(
    ["参数", "什么意思", "怎么调"],
    [
        ["语义回复阈值", "Camille 多大把握才回复。0.75 = 75%以上把握才回",
         "回复太多 → 调高到 0.85\n回复太少 → 调低到 0.65"],
        ["冷却时间", "回复一个用户后，多少秒内不再回复同一个人",
         "太频繁 → 调大到 180秒\n不够及时 → 调小到 60秒"],
    ],
    [3.5, 6, 6.5]
)

add_tip("不想 Camille 每句都回？把「语义回复阈值」调到 0.85，「冷却时间」调到 180 秒，效果立竿见影。")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 七、策略管理
# ═══════════════════════════════════════════════════

add_heading_styled("七、策略管理 — 调整回复风格", 1)

add_para("策略管理控制的是 Camille 「怎么回复」，而不是「要不要回复」。")

add_para("位置：左边菜单 → 策略管理", bold=True, color=BLUE)

add_table(
    ["能调的东西", "什么意思", "怎么调"],
    [
        ["温度", "回复的创意程度。越低越严谨，越高越自由",
         "客服场景建议 0.3（严谨）"],
        ["回复长度", "回复最多多少字",
         "默认 512 一般够用"],
        ["启用/禁用", "可以关掉某个策略",
         "不需要的策略关掉即可"],
        ["回复概率", "不是每次都回复（比如设 0.5 = 一半概率回）",
         "一般保持 1.0（百分百回）"],
    ],
    [3.5, 6, 6.5]
)

add_tip("大多数时候不需要改策略管理。如果觉得回复太啰嗦，把温度调低就行。")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 八、人工转接
# ═══════════════════════════════════════════════════

add_heading_styled("八、人工转接 — 搞不定时叫人", 1)

add_para("当一个用户在群里短时间内反复问同一个问题（比如 10 分钟问了 3 遍），Camille 会自动 @ 人工客服来处理。")

add_para("位置：系统设置 → 人工客服转接", bold=True, color=BLUE)

add_heading_styled("三步开启", 2)
add_numbered("勾选「启用」", "①")
add_numbered("填客服的 Telegram 用户名（不要带 @）", "②")
add_numbered("点「保存转接配置」", "③")

add_heading_styled("触发条件", 2)
add_table(
    ["设置", "什么意思", "默认值"],
    [
        ["同一问题连发几次触发", "用户问同样的话问几遍后叫人", "3 次"],
        ["统计窗口", "在多长时间内计算重复次数", "10 分钟"],
        ["触发后冷却", "叫过人之后多久不会再叫", "5 分钟"],
    ],
    [5, 7, 4]
)

add_heading_styled("触发后会发生什么", 2)
add_bullet("Camille 正常回复用户后，在回复末尾追加一段话 @ 人工客服")
add_bullet("同时把用户的消息转发到客服的私聊里")
add_bullet("客服收到一条带按钮的消息，点了就能跳到群里那条消息")

add_heading_styled("值班模式", 2)
add_table(
    ["模式", "什么意思"],
    [
        ["始终可 @（默认）", "任何时候都会叫人工"],
        ["仅手动值班", "只有在后台手动勾了「值班中」才叫人"],
        ["仅工作时间段", "按设好的上下班时间来"],
    ],
    [5, 11]
)

add_tip("大多数人选「始终可 @」就够了。")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 九、常用操作速查表
# ═══════════════════════════════════════════════════

add_heading_styled("九、常用操作速查表", 1)

add_para("下面列出日常最常用的操作，方便快速查找。")

add_table(
    ["我想…", "去哪里", "怎么做"],
    [
        ["看 Camille 是否在线", "仪表盘",
         "看「Telegram 连接」和「AI 状态」"],
        ["加一条新的问答", "知识库 → 新增条目",
         "填标题、触发词、回答 → 保存"],
        ["改一条已有的回答", "知识库 → 找到那条 → 编辑",
         "改完 → 保存"],
        ["审核 AI 生成的回答", "智能学习",
         "看草稿 → 通过/拒绝/编辑"],
        ["让 Camille 少回复一些", "系统设置 → 回复逻辑管理",
         "调高「语义回复阈值」或调大「冷却时间」"],
        ["让 Camille 多回复一些", "系统设置 → 回复逻辑管理",
         "调低「语义回复阈值」或调小「冷却时间」"],
        ["开启人工转接", "系统设置 → 人工客服转接",
         "启用 → 填用户名 → 保存"],
        ["改 Camille 的说话风格", "系统设置 → AI 提示词",
         "编辑提示词 → 保存"],
        ["看 Camille 今天处理了多少消息", "仪表盘",
         "看「收到消息」和「已回复」数字"],
        ["查看操作记录", "操作记录",
         "可以看谁在后台做了什么改动"],
    ],
    [5, 5.5, 5.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 十、遇到问题怎么办
# ═══════════════════════════════════════════════════

add_heading_styled("十、遇到问题怎么办", 1)

add_table(
    ["问题", "可能的原因", "怎么解决"],
    [
        ["Camille 不回复消息了",
         "1. 程序停了\n2. Telegram 断线\n3. 回复逻辑太严格",
         "1. 检查仪表盘状态\n2. 联系技术重启\n3. 调低回复阈值"],
        ["Camille 回复太多、刷屏",
         "1. 回复阈值太低\n2. 冷却时间太短",
         "系统设置 → 回复逻辑管理\n调高阈值到 0.85，冷却调到 180秒"],
        ["Camille 回答错了",
         "知识库里没有这个问题，或者回答过时了",
         "去知识库加一条新的，或者改掉旧的回答"],
        ["用户反复问但没有叫人工",
         "1. 人工转接没开\n2. 客服用户名没填\n3. 在冷却期内",
         "检查系统设置 → 人工客服转接\n确保启用了并填了用户名"],
        ["后台打不开",
         "1. 服务器没开\n2. 端口被占了",
         "联系技术人员检查服务器"],
    ],
    [4, 5.5, 6.5]
)

doc.add_paragraph()
add_para("仍然解决不了？", bold=True, size=Pt(12))
add_bullet("打开后台 → 「实时日志」，看最新的错误信息")
add_bullet("把错误信息截图发给技术人员")
add_bullet("技术人员可以远程查看 logs/app.log 文件")

doc.add_paragraph()
doc.add_paragraph()

# ── 页脚 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Camille 智能客服操作手册 · 完 —")
run.font.size = Pt(10)
run.font.color.rgb = GRAY
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 保存 ──
out_path = os.path.join(os.path.dirname(__file__), "Camille智能客服操作手册.docx")
doc.save(out_path)
print(f"✅ 文档已生成: {out_path}")
