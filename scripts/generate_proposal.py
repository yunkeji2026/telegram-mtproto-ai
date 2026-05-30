"""
生成「多平台 AI 聊天机器人 — 合作项目方案」Word 文档
依赖：pip install python-docx
输出：docs/AI_ChatBot_Proposal.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

# ── 颜色常量 ──────────────────────────────────────────
PRIMARY   = RGBColor(0x14, 0x5D, 0xA0)   # 深蓝
SECONDARY = RGBColor(0x00, 0xB4, 0xD8)   # 亮蓝
ACCENT    = RGBColor(0x06, 0xD6, 0xA0)   # 绿松石
DARK      = RGBColor(0x1A, 0x1A, 0x2E)   # 深黑
GRAY      = RGBColor(0x55, 0x55, 0x66)   # 正文灰
LIGHT_BG  = RGBColor(0xF0, 0xF7, 0xFF)   # 浅蓝背景（表格）
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

# ── 辅助函数 ──────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """设置表格单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(table):
    """给表格加细边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'D0E8FF')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font='微软雅黑'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1, color=None, size=None, center=False):
    sizes = {1: 22, 2: 16, 3: 13}
    colors = {1: PRIMARY, 2: PRIMARY, 3: DARK}
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size or sizes.get(level, 12))
    run.font.color.rgb = color or colors.get(level, DARK)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    pf = para.paragraph_format
    pf.space_before = Pt({1:18, 2:14, 3:10}.get(level, 8))
    pf.space_after  = Pt({1:8,  2:6,  3:4 }.get(level, 4))
    return para


def add_body(doc, text, color=None, size=11, indent=0, space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(indent)
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.line_spacing = Pt(20)
    add_run(para, text, size=size, color=color or GRAY)
    return para


def add_bullet(doc, text, bold_prefix=None, indent=0.6):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent  = Cm(indent)
    para.paragraph_format.space_after  = Pt(3)
    para.paragraph_format.line_spacing = Pt(18)
    if bold_prefix:
        add_run(para, bold_prefix, bold=True, size=11, color=PRIMARY)
        add_run(para, text, size=11, color=GRAY)
    else:
        add_run(para, text, size=11, color=GRAY)
    return para


def add_divider(doc, color='145DA0'):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def add_colored_box(doc, title, body_lines, bg='EBF5FF', title_color=None):
    """带背景色的信息框（用1×1表格模拟）"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cell.width = Inches(6.2)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, title + '\n', bold=True, size=11,
            color=title_color or PRIMARY)
    for line in body_lines:
        add_run(p, line + '\n', size=10, color=GRAY)
    doc.add_paragraph()   # spacing


def make_two_col_table(doc, rows_data, header=None, col_widths=(3.0, 3.5)):
    """制作两列对比/列表表格"""
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(tbl)
    if header:
        row = tbl.add_row()
        for i, h in enumerate(header):
            c = row.cells[i]
            set_cell_bg(c, '145DA0')
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, h, bold=True, size=11, color=WHITE)
    for item in rows_data:
        row = tbl.add_row()
        for i, val in enumerate(item):
            c = row.cells[i]
            set_cell_bg(c, 'F0F7FF' if rows_data.index(item) % 2 == 0 else 'FFFFFF')
            add_run(c.paragraphs[0], val, size=10.5, color=DARK)
    doc.add_paragraph()
    return tbl


def make_feature_table(doc, features):
    """功能特性三列表格: 图标/功能/说明"""
    tbl = doc.add_table(rows=0, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(tbl)
    # header
    hrow = tbl.add_row()
    for i, h in enumerate(['模块', '功能名称', '功能说明']):
        c = hrow.cells[i]
        set_cell_bg(c, '145DA0')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=11, color=WHITE)
    for idx, (mod, name, desc) in enumerate(features):
        row = tbl.add_row()
        bg = 'F0F7FF' if idx % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        set_cell_bg(row.cells[2], bg)
        add_run(row.cells[0].paragraphs[0], mod,  bold=True, size=10, color=SECONDARY)
        add_run(row.cells[1].paragraphs[0], name, bold=True, size=10, color=DARK)
        add_run(row.cells[2].paragraphs[0], desc, size=10,   color=GRAY)
    doc.add_paragraph()
    return tbl


# ══════════════════════════════════════════════════════
# 正文构建
# ══════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── 页面边距 ──
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ══════════════════ 封面 ══════════════════════════
    doc.add_paragraph('\n\n')

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover_title.add_run('多平台 AI 聊天机器人系统')
    r.bold = True
    r.font.size = Pt(32)
    r.font.color.rgb = PRIMARY
    r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = cover_sub.add_run('智能客服 · 全渠道触达 · 自动化引流')
    rs.font.size = Pt(16)
    rs.font.color.rgb = SECONDARY
    rs.font.name = '微软雅黑'
    rs._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph('\n')
    add_divider(doc, '00B4D8')
    doc.add_paragraph('\n')

    cover_tag = doc.add_paragraph()
    cover_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = cover_tag.add_run('合作项目方案书')
    rt.font.size = Pt(13)
    rt.font.color.rgb = GRAY
    rt.font.name = '微软雅黑'
    rt._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = date_para.add_run(datetime.date.today().strftime('%Y 年 %m 月'))
    rd.font.size = Pt(12)
    rd.font.color.rgb = GRAY
    rd.font.name = '微软雅黑'
    rd._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()

    # ══════════════════ 第一章：项目概述 ══════════════
    add_heading(doc, '一、项目概述', level=1)
    add_divider(doc)
    add_body(doc,
        '随着社交媒体与即时通讯平台的高速普及，企业面临的客户触点呈现出'
        '碎片化、多平台并存的格局。传统单一渠道的人工客服模式已无法满足'
        '7×24 小时全球用户的即时响应需求，同时也带来了高昂的人力成本与'
        '不一致的服务体验。')
    add_body(doc,
        '本方案所介绍的「多平台 AI 聊天机器人系统」，是一套融合大语言模型'
        '（LLM）、知识库检索、拟人化交互、以及跨平台 RPA 自动化技术的'
        '企业级智能客服解决方案。系统当前已稳定运行于 Telegram、LINE、'
        'Facebook Messenger 及 WhatsApp 四大主流平台，具备单系统同时'
        '管理多账号、多平台、多角色人设的能力，并配有完整的 Web 管理后台'
        '与可观测性监控体系。')

    add_heading(doc, '核心价值主张', level=2)
    make_two_col_table(doc, [
        ('降低运营成本',   '7×24 自动值守，减少 80%+ 人工重复回复'),
        ('提升响应速度',   '毫秒级 AI 推理，平均响应时间 < 3 秒'),
        ('多平台统一管理', '一套后台管理 Telegram / LINE / WA / Messenger'),
        ('人设定制化',     '支持多角色人设，匹配不同业务场景与语言风格'),
        ('数据驱动决策',   '完整漏斗、亲密度、审计日志，可视化运营看板'),
        ('快速行业落地',   '内置 9 大行业域包，开箱即用，最快 1 天上线'),
    ], header=['价值维度', '核心收益'])

    # ══════════════════ 第二章：支持平台 ══════════════
    add_heading(doc, '二、支持平台与接入方式', level=1)
    add_divider(doc)
    add_body(doc,
        '系统采用统一的 AI 内核（SkillManager + AIClient），每个平台独立'
        '一套 Runner 适配层，做到接入方式灵活、互不干扰，同时共享人设、'
        '知识库与用户记忆。')

    platforms = [
        ('Telegram',           'MTProto 原生协议',
         '直连官方 API，支持私聊/群组/频道；语音消息 ASR 转写 + TTS 语音回复；'
         '多账号并行；人工转接（@客服）；消息去重/限流'),
        ('LINE',               'Messaging API + 个人号 RPA',
         '官方 Webhook 接入（企业账号）；ADB+UIAutomator 个人号 RPA；'
         '多会话自动导航；分条回复；图片 OCR/Vision 兜底；审批模式'),
        ('Facebook Messenger', 'Android ADB RPA',
         'ADB 驱动手机 App；Vision AI 自动识别聊天人；坐标精准定位；'
         'Message Requests 自动接收；语音消息捕获与 ASR；引流到 LINE'),
        ('WhatsApp',           'Android ADB RPA',
         '直接读取文件系统 .opus 语音（零延迟）；TTS 通过系统分享发送；'
         'Quote Reply 引用回复；多媒体 Vision 理解；多账号同时运行'),
    ]

    tbl = doc.add_table(rows=0, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(tbl)
    hrow = tbl.add_row()
    for i, h in enumerate(['平台', '接入技术', '主要能力']):
        c = hrow.cells[i]
        set_cell_bg(c, '145DA0')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=11, color=WHITE)
    for idx, (plat, tech, cap) in enumerate(platforms):
        row = tbl.add_row()
        bg = 'EBF5FF' if idx % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        set_cell_bg(row.cells[2], bg)
        add_run(row.cells[0].paragraphs[0], plat, bold=True, size=11, color=PRIMARY)
        add_run(row.cells[1].paragraphs[0], tech, size=10,   color=SECONDARY)
        add_run(row.cells[2].paragraphs[0], cap,  size=10,   color=GRAY)
    doc.add_paragraph()

    # ══════════════════ 第三章：核心功能 ══════════════
    add_heading(doc, '三、核心功能详解', level=1)
    add_divider(doc)

    # 3.1 AI 回复引擎
    add_heading(doc, '3.1  AI 回复引擎（SkillManager）', level=2)
    add_body(doc,
        'AI 回复引擎是系统的核心大脑，采用四层触发-执行架构，'
        '确保每一条消息都能得到精准、自然的响应。')
    make_feature_table(doc, [
        ('意图识别',   '四层意图触发',    '关键词 → 正则 → BM25 知识库 → LLM 大模型，逐层降级，兼顾速度与准确性'),
        ('知识库检索', 'BM25 混合检索',   '领域知识库支持中文分词 + BM25 算法精准召回，结果注入 LLM 上下文，确保回答有据可查'),
        ('上下文管理', '多轮对话记忆',    '保留最近 16 轮对话历史，支持情景记忆（EpisodicMemory）跨会话个性化'),
        ('人设系统',   '多角色人设池',    '支持暖陪聊（warm_companion）/ 专业客服（professional_support）等多人设，按用户/场景动态切换'),
        ('语言守卫',   '自动语言跟随',    '自动检测用户语言（中/英/日/繁等），AI 回复自动切换，无需配置'),
        ('回复策略',   '五级回复策略',    'S1-S5 策略调度：从主动回复到静默概率控制，避免过度打扰'),
        ('冷却限流',   '多维冷却机制',    '用户/内容/全局三维冷却，防止重复/骚扰回复；daily_cap 每日发送上限'),
        ('情感引擎',   'EmotionEnhancer',  '分析用户情绪，在亲密关系不同阶段注入对应情感指令，回复更有温度'),
    ])

    # 3.2 语音能力
    add_heading(doc, '3.2  语音能力（ASR + TTS）', level=2)
    add_body(doc,
        '系统内置完整的语音输入（识别）与语音输出（合成）管线，支持多引擎'
        '热切换与自动降级，满足语音交互场景需求。')
    add_bullet(doc, 'Whisper Local / Faster-Whisper / OpenAI Whisper API — 三引擎自动降级', bold_prefix='ASR 识别：')
    add_bullet(doc, '支持 .ogg / .opus / .mp3 等多格式语音文件，自动解码转写', bold_prefix='格式兼容：')
    add_bullet(doc, 'Edge TTS（微软 Azure）/ OpenAI TTS / 声音克隆命令行 — 灵活配置', bold_prefix='TTS 合成：')
    add_bullet(doc, '可注入本人声音样本，实现高保真声音克隆回复', bold_prefix='声音克隆：')
    add_bullet(doc, '对方发语音 → 自动触发语音回复；或随机概率/强制语音模式', bold_prefix='触发策略：')
    add_bullet(doc, '合成超时/超长 → 自动回退文字；发送失败 → BACK+HOME 状态恢复', bold_prefix='容错降级：')
    doc.add_paragraph()

    # 3.3 多媒体理解
    add_heading(doc, '3.3  多媒体消息理解', level=2)
    add_body(doc,
        '系统支持图片、语音、视频、贴纸等多种消息类型的智能理解，'
        '通过 Vision AI 与 OCR 双通道确保内容不丢失。')
    make_two_col_table(doc, [
        ('图片理解', '截图裁剪 → GLM-4V / Ollama 多模态描述 → 注入 AI 上下文'),
        ('语音转写', '自动识别语音气泡 → 下载 → Whisper 转文字 → AI 回复'),
        ('OCR 兜底', '视觉 AI 失效时走 Tesseract OCR，保证文字内容可读'),
        ('贴纸/GIF', 'Vision 识别贴纸含义，AI 理解表情语义后自然回复'),
        ('文件消息', '文件类型识别 + 占位提示，引导用户改发文字说明'),
    ], header=['媒体类型', '处理流程'])

    # 3.4 拟人化交互
    add_heading(doc, '3.4  拟人化交互设计', level=2)
    add_body(doc,
        '系统深度模拟真人聊天节奏，有效规避平台风控，同时大幅提升用户体验。')
    add_bullet(doc, '阅读停顿 0.8–2.0 秒，打字延时按字数动态计算，每字 40–80ms', bold_prefix='阅读 & 打字节奏：')
    add_bullet(doc, '长回复自动按句子/长度分拆成多条发送，每条间隔 700–1800ms', bold_prefix='分条发送：')
    add_bullet(doc, '检测到对方连续发消息时，等待"安静窗口"再生成回复，避免打断', bold_prefix='安静窗口：')
    add_bullet(doc, '对方多条消息自动合并分析意图，支持 Quote Reply 精准引用回复', bold_prefix='多条消息聚合：')
    add_bullet(doc, '所有停顿加 ±25% 随机抖动，避免机械规律性，更像真人', bold_prefix='随机抖动：')
    doc.add_paragraph()

    # 3.5 人工转接
    add_heading(doc, '3.5  人工客服转接（Human Escalation）', level=2)
    add_body(doc,
        '当 AI 无法解决用户问题或用户重复询问同类问题时，系统自动触发'
        '人工客服 @ 通知，确保关键客户需求不漏接。')
    add_bullet(doc, '同一问题重复 N 次（可配置）自动触发，避免骚扰性转接')
    add_bullet(doc, '支持排班计划：按工作时间/班次决定是否触发转接')
    add_bullet(doc, '支持多客服团队轮询分配，消息附带原始问句链接，方便跳转')
    add_bullet(doc, '转接后向各客服私发原始消息转发，内联按钮直达群内原话')
    add_bullet(doc, '120 秒冷却去重，防止同一问题多次触发转接')
    doc.add_paragraph()

    # 3.6 Contacts/Handoff 引流子系统
    add_heading(doc, '3.6  跨平台引流子系统（Contacts / Handoff）', level=2)
    add_body(doc,
        '这是系统独有的跨平台用户生命周期管理模块，专为「Messenger 引流 → '
        'LINE 私域」等跨平台转化链路设计。')
    make_feature_table(doc, [
        ('Contact 管理',    '联系人全生命周期',  '统一记录跨平台联系人，Journey 状态机跟踪每个用户从初识到转化'),
        ('HandoffToken',    '引流暗号机制',      '生成唯一引流码，用户添加 LINE 时自动识别来源渠道，归因精准'),
        ('Readiness 引擎',  '智能引流时机',      '基于对话轮数 + 亲密度评分计算 readiness（0-100），高于阈值才触发引流'),
        ('Reactivation',    '沉默用户唤醒',      '检测沉默 N 天以上的老用户，自动调度唤醒话术，回流率提升 3 倍以上'),
        ('漏斗看板',        '可视化转化漏斗',    'Web 后台实时展示各阶段用户数、转化率、每日引流量趋势'),
        ('daily_cap 限控',  '合规发送上限',      '每账号每日最多 N 次引流，全局上限可配，防止账号被封'),
    ])

    # ══════════════════ 第四章：行业域包 ══════════════
    add_heading(doc, '四、行业域包（开箱即用）', level=1)
    add_divider(doc)
    add_body(doc,
        '系统内置 9 大行业域包，每个域包包含：系统提示词、知识库种子、'
        '话术模板、意图关键词、i18n 多语言支持与 Web 扩展页面。'
        '合作伙伴可直接选用或在此基础上快速定制。')

    domains_data = [
        ('payment',      '支付通道客服',   '订单查询、GXP 指令代发、通道额度管理、费率查询、通道健康监控仪表盘'),
        ('ecommerce',    '电商客服',       '商品咨询、订单跟踪、退换货、促销活动、智能推荐'),
        ('crypto',       '加密货币行业',   '币价查询、交易所咨询、出入金指导、合规提示、市场行情播报'),
        ('conversion',   '销售转化/陪聊',  '暖场破冰、关系阶段推进（initial→warming→intimate→steady）、情感管理'),
        ('community',    '社区运营',       '群公告管理、活动通知、新人欢迎、积分查询、用户分层维护'),
        ('education',    '在线教育',       '课程咨询、报名引导、学习进度查询、作业提醒、课后答疑'),
        ('it_helpdesk',  'IT 技术支持',    '故障报告、工单创建、已知问题库、版本更新通知、SLA 跟踪'),
        ('legal',        '法律咨询',       '常见法律问题 FAQ、条款解释、免责声明、专业律师转接'),
        ('general',      '通用助手',       '适用任意场景的基础知识问答与闲聊陪伴底座'),
    ]

    tbl = doc.add_table(rows=0, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(tbl)
    hrow = tbl.add_row()
    for h in ['域包 ID', '适用行业', '核心功能覆盖']:
        c = hrow.cells[['域包 ID', '适用行业', '核心功能覆盖'].index(h)]
        set_cell_bg(c, '145DA0')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=11, color=WHITE)
    for idx, (did, dname, ddesc) in enumerate(domains_data):
        row = tbl.add_row()
        bg = 'EBF5FF' if idx % 2 == 0 else 'FFFFFF'
        for i, (val, bold, col) in enumerate([
            (did, True, SECONDARY), (dname, True, DARK), (ddesc, False, GRAY)
        ]):
            set_cell_bg(row.cells[i], bg)
            add_run(row.cells[i].paragraphs[0], val, bold=bold, size=10, color=col)
    doc.add_paragraph()

    add_colored_box(doc,
        '💡 定制域包服务',
        [
            '我们提供域包快速定制服务，根据贵方行业知识库、话术风格与业务流程，',
            '通常 3-5 个工作日即可完成定制域包交付，直接对接上线运行。',
        ],
        bg='E8F8F5', title_color=ACCENT)

    # ══════════════════ 第五章：技术亮点 ══════════════
    add_heading(doc, '五、技术亮点与系统优势', level=1)
    add_divider(doc)

    add_heading(doc, '5.1  多设备自动化管理', level=2)
    add_body(doc,
        '系统搭载 Device Coordinator + HotPlug Watcher，支持多台 Android '
        '手机同时运行，设备插拔自动识别纳管，无需人工干预。')
    add_bullet(doc, '支持 4+ 台手机同时运行，每台可承载 1–3 个平台账号')
    add_bullet(doc, 'ADB 热插拔自动识别，新设备插入 15s 内自动激活')
    add_bullet(doc, '设备故障熔断机制：连续失败触发自动暂停 + Webhook 告警')
    add_bullet(doc, '设备统计面板：实时查看每台设备状态、发送量、错误率')
    doc.add_paragraph()

    add_heading(doc, '5.2  全链路可观测性', level=2)
    add_body(doc,
        '企业级监控体系，确保每一个运营动作都有据可查、异常可感知。')
    make_two_col_table(doc, [
        ('Prometheus 指标',   '消息处理量/API 调用/响应时间/错误率，Grafana 仪表盘可视化'),
        ('结构化审计日志',    '每条发送记录 JSON 落盘，含时间戳/账号/平台/内容哈希'),
        ('Webhook 告警',      '熔断/设备掉线/配置变更 → 即时推送到 Telegram Bot 或 Slack'),
        ('失败截图留痕',      '发送失败时自动截图保存，Web 后台可查看复现故障'),
        ('健康检查对账',      '定时 dumpsys 对账，连续疑似漏读触发告警'),
        ('指标 JSONL 记录',   '每轮 run_once 结果结构化写入，支持外部日志分析系统接入'),
    ], header=['监控维度', '实现说明'])

    add_heading(doc, '5.3  安全与合规设计', level=2)
    add_bullet(doc, '所有配置项通过 YAML 文件管理，API Key 不入代码仓库')
    add_bullet(doc, '请求体大小限制（2MB 默认）+ 限流 Token Bucket，防爆破攻击')
    add_bullet(doc, '反代真实 IP 解析，防 XFF 伪造绕过限流')
    add_bullet(doc, 'Web 后台 RBAC 角色权限（master / admin / viewer）分级控制')
    add_bullet(doc, '声音克隆需配置 owner_consent: true，杜绝非授权克隆')
    add_bullet(doc, '每日发送上限 + 账号冷却，防平台封号风险')
    doc.add_paragraph()

    add_heading(doc, '5.4  AI 提供商灵活切换', level=2)
    add_body(doc,
        '系统 AI 层完全解耦，支持一行配置切换大模型提供商，'
        '不锁定任何单一 AI 服务，保障供应商议价能力与业务连续性。')
    make_two_col_table(doc, [
        ('DeepSeek',       'deepseek-chat / deepseek-reasoner，成本极低，中文能力强'),
        ('OpenAI',         'GPT-4o / GPT-4o-mini，旗舰模型，英文+复杂推理首选'),
        ('Ollama 本地',    '私有化部署，数据不出内网，适合高合规要求场景'),
        ('其他兼容接口',   '任意 OpenAI-compatible API（Claude / Gemini 代理等）均可接入'),
    ], header=['AI 提供商', '适用说明'])

    # ══════════════════ 第六章：Web 后台 ══════════════
    add_heading(doc, '六、Web 管理后台', level=1)
    add_divider(doc)
    add_body(doc,
        '系统配套完整的 Web 管理后台（FastAPI + Jinja2 模板），'
        '无需命令行即可完成日常运营管理。默认端口 18787，支持多角色登录。')

    add_heading(doc, '主要功能页面', level=2)
    make_feature_table(doc, [
        ('/telegram',       'Telegram 设置',      '账号信息、今日统计、语音回复设置、ASR 配置、日志流实时查看'),
        ('/line-rpa',       'LINE RPA 控制台',     '服务启停、轮询配置、失败截图库、健康检查状态、审批队列'),
        ('/messenger-rpa',  'Messenger RPA 控制台','多账号管理、对话配置、语音设置、转化配置、运营指标'),
        ('/rpa-overview',   '跨平台总览',          '所有平台/账号实时状态、intent_tags 关键词可视化编辑'),
        ('/personas',       '人设管理 Studio',     '人设创建/编辑、语言风格调试、TTS 试听、实时预览'),
        ('/episodic-memory','情景记忆管理',         '用户记忆查询、条目审核、手动注入关键记忆片段'),
        ('/channels',       '通道管理（支付域）',  '实时通道健康度、额度看板、费率管理、通道告警'),
        ('/ops/contacts',   '联系人漏斗',          'Contacts 全列表、Journey 状态、引流统计、合并去重'),
        ('/admin',          '系统管理',            '用户管理、配置热更新、审计日志查看、Webhook 配置'),
    ])

    # ══════════════════ 第七章：适用场景 ══════════════
    add_heading(doc, '七、典型应用场景', level=1)
    add_divider(doc)

    scenarios = [
        ('跨境电商自动客服',
         '覆盖 Telegram 群组 + WhatsApp 私聊双渠道，AI 自动处理商品咨询、'
         '物流查询、售后退款等高频问题，人工客服只需处理 10% 的复杂工单，'
         '整体客服成本降低 75%。'),
        ('社交电商引流转化',
         '在 Facebook Messenger 建立初步联系后，系统自动在合适时机'
         '（readiness ≥ 70）发送 LINE 引流暗号，实现从公域流量到私域'
         '沉淀的自动化闭环。亲密度引擎确保引流话术自然不突兀。'),
        ('加密货币/支付行业',
         '内置 crypto 与 payment 域包，支持通道状态查询、GXP 指令代发、'
         '实时汇率播报。窄回复模式（narrow_reply）确保非业务话题不占用'
         '客服资源，符合行业合规要求。'),
        ('在线教育招生咨询',
         '多语言自动识别，覆盖中英日韩等学员群体；知识库快速配置课程信息；'
         '感兴趣用户自动进入报名引导流程；人工教务老师在关键节点介入。'),
        ('社群运营与会员服务',
         '自动欢迎新成员、回答社群规则、推送活动通知；'
         '沉默会员唤醒机制（Reactivation）定期发送个性化关怀消息，'
         '社群活跃度提升显著。'),
        ('本地服务/餐饮/零售',
         '营业时间外自动接待，收集用户需求和联系方式；'
         '工作时间内消息自动转人工，附带对话摘要，员工可即时上手处理。'),
    ]

    for title, desc in scenarios:
        add_heading(doc, title, level=2, size=13)
        add_body(doc, desc)
        add_divider(doc, 'D0E8FF')

    # ══════════════════ 第八章：合作方案 ══════════════
    add_heading(doc, '八、合作项目方案', level=1)
    add_divider(doc)
    add_body(doc,
        '我们提供灵活多样的合作模式，无论是快速试用、定制交付还是'
        '深度战略合作，均可根据贵方实际需求量身定制。')

    add_heading(doc, '8.1  合作模式', level=2)
    make_two_col_table(doc, [
        ('模式 A：SaaS 授权',
         '按账号/月计费；系统部署在我方云端；贵方通过 Web 后台管理；'
         '适合快速启动、无需自建运维的中小团队'),
        ('模式 B：私有化部署',
         '源码授权 + 一次性交付；部署在贵方服务器/内网；'
         '数据完全自主；适合金融/支付等高合规要求场景'),
        ('模式 C：联合开发',
         '以本系统为底座，共同开发行业特化版本；'
         '收益按约定比例分成；适合拥有行业资源的垂直领域合作伙伴'),
        ('模式 D：集成服务',
         '提供 API/SDK 接入现有产品；系统作为 AI 回复中台；'
         '适合已有用户体系希望快速添加 AI 能力的产品团队'),
    ], header=['合作模式', '适用场景与说明'])

    add_heading(doc, '8.2  标准交付物', level=2)
    add_bullet(doc, '完整系统源码（含 9 大域包 + 测试套件 300+ 用例）')
    add_bullet(doc, '一键部署脚本 + Docker Compose HA 高可用配置')
    add_bullet(doc, 'Web 管理后台 + Grafana 监控仪表盘')
    add_bullet(doc, '配置手册 + 运维操作手册（中文版）')
    add_bullet(doc, '域包快速定制服务（按行业需求 3-5 个工作日）')
    add_bullet(doc, '上线后 30 天技术支持（问题响应时间 < 4 小时）')
    doc.add_paragraph()

    add_heading(doc, '8.3  标准实施周期', level=2)
    impl_steps = [
        ('第 1 周', '需求对接 + 域包定制',
         '了解业务场景、知识库整理、话术确认、AI 人设调优'),
        ('第 2 周', '环境部署 + 平台接入',
         '服务器/设备准备、平台账号接入、基础功能联调'),
        ('第 3 周', '灰度上线 + 数据调优',
         '小规模上线、意图识别优化、回复质量审核与迭代'),
        ('第 4 周', '全量上线 + 交付',
         '全量开放、监控配置、团队培训、文档交付'),
    ]
    tbl = doc.add_table(rows=0, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(tbl)
    hrow = tbl.add_row()
    for h in ['阶段', '里程碑', '主要工作内容']:
        i = ['阶段', '里程碑', '主要工作内容'].index(h)
        c = hrow.cells[i]
        set_cell_bg(c, '145DA0')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=11, color=WHITE)
    for idx, (phase, milestone, work) in enumerate(impl_steps):
        row = tbl.add_row()
        bg = 'EBF5FF' if idx % 2 == 0 else 'FFFFFF'
        for i, (val, bold, col) in enumerate([
            (phase, True, PRIMARY), (milestone, True, DARK), (work, False, GRAY)
        ]):
            set_cell_bg(row.cells[i], bg)
            add_run(row.cells[i].paragraphs[0], val, bold=bold, size=10, color=col)
    doc.add_paragraph()

    # ══════════════════ 第九章：未来规划 ══════════════
    add_heading(doc, '九、未来发展规划', level=1)
    add_divider(doc)
    add_body(doc,
        '基于现有稳定架构，我们规划了清晰的产品演进路线，'
        '确保系统持续领先行业水平，为合作伙伴创造长期价值。')

    roadmap = [
        ('2026 Q3',  '🤝 跨平台身份统一（CrossPlatformIdentity）',
         '同一用户在 Telegram / LINE / WA / Messenger 的记忆贯通，'
         '实现真正的"认识你"而非"认识你的账号"'),
        ('2026 Q3',  '🎭 关系阶段引擎（RelationshipStager）',
         '将亲密度评分 0-100 映射到 4 个关系阶段，'
         'AI 提示词随阶段动态调整，从陌生人到老朋友的完整体验'),
        ('2026 Q4',  '🎛️ AI Studio 一体化中台',
         '人设/记忆/知识审核/关系看板四合一集中入口，'
         '运营人员无需懂代码即可完成全部 AI 行为调优'),
        ('2026 Q4',  '📊 智能运营分析',
         '对话质量评分、意图分布热力图、用户留存预测、'
         'A/B 回复策略自动对比测试'),
        ('2027 Q1',  '🌐 Web 端插件化',
         '支持网站 / 小程序 / APP 端 AI 客服插件，'
         '统一接入 AI 内核，扩展到 WhatsApp Web / Instagram 等新平台'),
        ('2027 Q1',  '🧠 垂直领域大模型微调',
         '针对支付/电商/法律等高频行业，利用积累的对话数据'
         '进行领域微调，显著提升专业问答准确率'),
    ]

    tbl = doc.add_table(rows=0, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(tbl)
    hrow = tbl.add_row()
    for h in ['时间节点', '功能方向', '价值说明']:
        i = ['时间节点', '功能方向', '价值说明'].index(h)
        c = hrow.cells[i]
        set_cell_bg(c, '06D6A0')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=11, color=WHITE)
    for idx, (ts, feat, val) in enumerate(roadmap):
        row = tbl.add_row()
        bg = 'EAFAF5' if idx % 2 == 0 else 'FFFFFF'
        for i, (v, bold, col) in enumerate([
            (ts, True, ACCENT), (feat, True, DARK), (val, False, GRAY)
        ]):
            set_cell_bg(row.cells[i], bg)
            add_run(row.cells[i].paragraphs[0], v, bold=bold, size=10, color=col)
    doc.add_paragraph()

    # ══════════════════ 第十章：下一步 ══════════════
    add_heading(doc, '十、下一步行动', level=1)
    add_divider(doc)

    add_colored_box(doc,
        '🚀  立即开始合作',
        [
            '我们邀请贵方进行一次 60 分钟的在线产品演示，',
            '届时将展示系统实时运行效果，并针对贵方业务场景进行方案定制讨论。',
            '',
            '演示内容：',
            '  · AI 回复全链路实时展示（从接收消息到发出回复）',
            '  · Web 管理后台操作演示',
            '  · 域包切换与人设调优演示',
            '  · 多平台并发运行演示',
            '',
            '请通过以下方式与我们联系，安排演示时间：',
            '  · Telegram：@your_contact',
            '  · 邮箱：contact@yourcompany.com',
            '  · 微信：your_wechat_id',
        ],
        bg='E8F8F5', title_color=ACCENT)

    add_body(doc,
        '我们期待与贵方建立长期合作关系，共同在 AI 智能客服领域创造更大价值。',
        color=DARK, size=12)

    doc.add_paragraph('\n')
    add_divider(doc, '00B4D8')

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_para,
        f'本文档生成于 {datetime.date.today().strftime("%Y 年 %m 月 %d 日")}  ·  '
        '多平台 AI 聊天机器人系统  ·  内部保密文件',
        size=9, color=GRAY, italic=True)

    return doc


if __name__ == '__main__':
    out_dir = os.path.join(os.path.dirname(__file__), '..', 'docs')
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, 'AI_ChatBot_Proposal.docx')
    doc = build_document()
    doc.save(out_path)
    print(f'✅  文档已生成：{os.path.abspath(out_path)}')
